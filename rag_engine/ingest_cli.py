#!/usr/bin/env python3
"""
CLI script for PDF/DOCX ingestion - called from Node.js handler.
Extracts text, tables, and visual content (diagrams/charts via vision model),
chunks via rag_engine, outputs JSON.
"""
import sys
import os
import json
import io
import base64
import warnings

# Suppress deprecation warnings BEFORE any library imports so they
# don't contaminate stdout (which must be pure JSON for the Node handler).
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

# Add rag_engine to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__)))

from chunking import chunk_text

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Prefer 'pymupdf' (new import name, no deprecation warning).
# Fall back to 'fitz' (old name, prints a warning to stdout on newer versions).
# Redirect stdout during import so any warning print can't contaminate JSON output.
HAS_PYMUPDF = False
fitz = None

for _mod_name in ("pymupdf", "fitz"):
    try:
        _saved_stdout = sys.stdout
        sys.stdout = io.StringIO()
        fitz = __import__(_mod_name)
        sys.stdout = _saved_stdout
        HAS_PYMUPDF = True
        break
    except ImportError:
        sys.stdout = _saved_stdout
        continue

try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
    # Set Tesseract path for Windows
    if os.name == 'nt':
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Users\Parth\AppData\Local\Programs\Tesseract-OCR\tesseract.exe',
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"Tesseract found at: {path}", file=sys.stderr)
                break
        else:
            print("WARNING: Tesseract executable not found in standard locations", file=sys.stderr)
except ImportError:
    HAS_TESSERACT = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False


def _get_gemini_key():
    """Get Gemini API key from environment."""
    k = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or os.getenv("GEMMA_API_KEY")
    if k and len(k.strip()) >= 10:
        return k.strip()
    return None


def _get_gemini_model():
    """Get the vision-capable Gemini model."""
    return os.getenv("GEMINI_MODEL", "gemini-1.5-flash")


def describe_image_with_vision(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """
    Use Gemini's vision API to describe an image (diagram, chart, flowchart, etc.)
    Returns a structured text description that can be chunked alongside regular text.
    """
    key = _get_gemini_key()
    if not key or not HAS_REQUESTS:
        return ""

    model = _get_gemini_model()
    # Use a vision-capable model — gemini-1.5-flash and gemini-2.0-flash both support vision
    vision_models = [model, "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]
    # Deduplicate while preserving order
    seen = set()
    vision_models = [m for m in vision_models if not (m in seen or seen.add(m))]

    b64 = base64.b64encode(image_bytes).decode("utf-8")

    prompt = (
        "You are analyzing a diagram, chart, flowchart, or figure from an educational document. "
        "Provide a detailed text description that captures ALL the information shown. "
        "Include:\n"
        "- The type of visual (flowchart, bar chart, diagram, table, etc.)\n"
        "- All labels, nodes, and text elements visible\n"
        "- The structure and relationships (arrows, connections, hierarchy)\n"
        "- Any data values, axes, or numerical information\n"
        "- The overall meaning/conclusion the visual conveys\n\n"
        "Format as clear, structured text. Be thorough — this description will be used "
        "for question generation, so capture every detail."
    )

    for m in vision_models:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{m}:generateContent?key={key}"
            payload = {
                "contents": [{"parts": [
                    {"text": prompt},
                    {"inlineData": {"mimeType": mime_type, "data": b64}}
                ]}],
                "generationConfig": {"temperature": 0.4, "maxOutputTokens": 2048}
            }
            r = requests.post(url, json=payload, timeout=60)
            if r.status_code == 404:
                continue
            r.raise_for_status()
            j = r.json()
            parts = j.get("candidates", [{}])[0].get("content", {}).get("parts", [])
            for part in parts:
                if part.get("text") and not part.get("thought"):
                    return part["text"]
            if parts and parts[-1].get("text"):
                return parts[-1]["text"]
        except Exception as e:
            print(f"Vision model error ({m}): {e}", file=sys.stderr)
            continue

    return ""


def table_to_markdown(table: list) -> str:
    """Convert a pdfplumber table (list of rows) to Markdown format."""
    if not table or not table[0]:
        return ""
    # Clean cells
    rows = []
    for row in table:
        cleaned = [(cell or "").strip() if cell else "" for cell in row]
        rows.append(cleaned)

    # Build markdown table
    lines = []
    # Header
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row to match header length
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract text from PDF using pdfplumber (with table extraction),
    then PyMuPDF, then OCR fallback. Also uses Gemini vision for
    diagrams/charts/flowcharts on each page.
    """
    text_parts = []
    tables_found = 0
    images_described = 0

    # ── Pass 1: pdfplumber — text + tables ──
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract regular text
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)

                    # Extract tables
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            if table and len(table) > 1:
                                md = table_to_markdown(table)
                                if md:
                                    text_parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")
                                    tables_found += 1
                    except Exception as te:
                        print(f"Table extraction error on page {page_num + 1}: {te}", file=sys.stderr)

            if text_parts:
                print(f"pdfplumber: extracted text + {tables_found} tables", file=sys.stderr)
                # Don't return yet — we still want to check for diagrams/charts
        except Exception as e:
            print(f"pdfplumber error: {e}", file=sys.stderr)

    # ── Pass 2: PyMuPDF — text (if pdfplumber got nothing) + vision for images ──
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(filepath)

            # If we got no text from pdfplumber, try PyMuPDF text extraction
            if not text_parts:
                for page in doc:
                    t = page.get_text()
                    if t:
                        text_parts.append(t)

            # ── Vision model: describe diagrams/charts/flowcharts on each page ──
            key = _get_gemini_key()
            if key and HAS_REQUESTS:
                for page_num, page in enumerate(doc):
                    # Get images on this page
                    images = page.get_images(full=True)
                    for img_info in images:
                        xref = img_info[0]
                        try:
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            mime = base_image.get("ext", "png")
                            mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff"}
                            mime_type = mime_map.get(mime, "image/png")

                            # Skip tiny icons (< 5KB) — likely not diagrams
                            if len(image_bytes) < 5000:
                                continue

                            print(f"Describing image on page {page_num + 1} ({len(image_bytes)} bytes) via vision model...", file=sys.stderr)
                            description = describe_image_with_vision(image_bytes, mime_type)
                            if description:
                                text_parts.append(f"\n[VISUAL CONTENT — Page {page_num + 1}]\n{description}\n[/VISUAL CONTENT]\n")
                                images_described += 1
                        except Exception as ve:
                            print(f"Vision extraction error on page {page_num + 1}: {ve}", file=sys.stderr)

                    # Also render the full page as an image and check for diagrams
                    # (only if the page has very little text — likely a diagram-heavy page)
                    page_text = page.get_text().strip()
                    if key and len(page_text) < 100 and not images:
                        try:
                            pix = page.get_pixmap(dpi=150)
                            img_data = pix.tobytes("png")
                            if len(img_data) > 5000:
                                print(f"Page {page_num + 1} has little text — rendering for vision model...", file=sys.stderr)
                                description = describe_image_with_vision(img_data, "image/png")
                                if description:
                                    text_parts.append(f"\n[VISUAL CONTENT — Page {page_num + 1}]\n{description}\n[/VISUAL CONTENT]\n")
                                    images_described += 1
                        except Exception as ve:
                            print(f"Page render vision error: {ve}", file=sys.stderr)

            doc.close()

            if text_parts:
                print(f"PyMuPDF: {images_described} images described via vision model", file=sys.stderr)
                if images_described > 0 or not text_parts:
                    pass  # Continue to OCR if needed
                else:
                    return "\n\n".join(text_parts)
        except Exception as e:
            print(f"PyMuPDF error: {e}", file=sys.stderr)

    # ── Pass 3: OCR fallback for scanned/image-based PDFs ──
    # Run OCR if we have very little meaningful text (likely a scanned PDF)
    total_text_len = sum(len(t) for t in text_parts)
    if total_text_len < 100 and HAS_PYMUPDF and HAS_TESSERACT:
        try:
            print(f"Little text extracted ({total_text_len} chars) — attempting OCR with Tesseract...", file=sys.stderr)
            doc = fitz.open(filepath)
            ocr_text_parts = []
            
            for page_num, page in enumerate(doc):
                # First try embedded images on this page
                images = page.get_images()
                page_ocr_text = ""
                
                for img in images:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image['image']
                    image = Image.open(io.BytesIO(image_bytes))
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        page_ocr_text += text + "\n"
                
                # If no text from embedded images, render page as image and OCR
                if not page_ocr_text.strip():
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        page_ocr_text += text + "\n"
                
                if page_ocr_text.strip():
                    ocr_text_parts.append(page_ocr_text)
                    print(f"Page {page_num + 1} OCR text length: {len(page_ocr_text)}", file=sys.stderr)

            doc.close()
            
            # If OCR got meaningful text, use it instead
            ocr_total = sum(len(t) for t in ocr_text_parts)
            if ocr_total > total_text_len:
                print(f"OCR extracted {ocr_total} chars (vs {total_text_len} from pdfplumber) — using OCR result", file=sys.stderr)
                text_parts = ocr_text_parts
            else:
                print(f"OCR extracted {ocr_total} chars — keeping original extraction", file=sys.stderr)
                
        except Exception as e:
            print(f"OCR error: {e}", file=sys.stderr)

    return "\n\n".join(text_parts) if text_parts else ""

def extract_text_from_docx(filepath: str) -> str:
    """Extract text and tables from DOCX using python-docx."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed")
    doc = docx.Document(filepath)
    text_parts = []

    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            md = table_to_markdown(rows)
            if md:
                text_parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")

    return "\n".join(text_parts)

def extract_text(filepath: str, filename: str) -> str:
    """Extract text based on file extension."""
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(filepath)
    elif ext in ('txt', 'md'):
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Usage: python ingest_cli.py <filepath> <filename>"}))
        sys.stdout.flush()
        sys.exit(1)
    
    filepath = sys.argv[1]
    filename = sys.argv[2]
    
    try:
        # Extract text (includes tables as markdown and visual descriptions from vision model)
        text = extract_text(filepath, filename)
        if not text or len(text.strip()) < 50:
            print(json.dumps({"error": "Could not extract meaningful text from file. The file may be a scanned/image PDF without OCR support, or the document may be empty."}))
            sys.stdout.flush()
            sys.exit(1)
        
        # Chunk using rag_engine
        title = filename.rsplit('.', 1)[0]
        chunks = chunk_text(text, title=title)

        # Count extracted content types
        table_count = text.count("[TABLE]")
        visual_count = text.count("[VISUAL CONTENT")
        
        # Output result
        result = {
            "title": title,
            "content": text,
            "chunks": chunks,
            "extraction_stats": {
                "tables": table_count,
                "visuals": visual_count,
                "total_chunks": len(chunks)
            }
        }
        print(json.dumps(result))
        sys.stdout.flush()
        
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.stdout.flush()
        sys.exit(1)

if __name__ == '__main__':
    main()